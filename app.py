import os
import fitz  # PyMuPDF pip install mymupdf
import asyncio
from typing import List
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
from chainlit.input_widget import Select, Slider
from langchain.docstore.document import Document
from langchain.memory import ConversationBufferMemory
from langchain_community.chat_message_histories import ChatMessageHistory
import chainlit as cl
from langchain_community.vectorstores.utils import filter_complex_metadata
class SentenceAwareTextSplitter(RecursiveCharacterTextSplitter):
    def __init__(self, chunk_size: int, chunk_overlap: int, separators: List[str] = None):
        super().__init__(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators if separators else [". ", "? ", "! ", "\n"]

    def split_text(self, text: str) -> List[str]:
        chunks = []
        current_chunk = ""
        for char in text:
            current_chunk += char
            if any(current_chunk.endswith(sep) for sep in self.separators):
                if len(current_chunk) >= self.chunk_size:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""

        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks
    
def extract_text_from_pdf(pdf_path: str) -> List[Document]:
    documents = []
    headers = []
    last_header = ""
    font_size_threshold = cl.user_session.get("Font_Size_Threshold", 11.9)
    with fitz.open(pdf_path) as pdf_doc:
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            text_dict = page.get_text("dict")
            page_text = ""
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["size"] > font_size_threshold and span["bbox"][2] - span["bbox"][0] > 30:
                                headers.append(span["text"])
                                last_header = span["text"]
                            page_text += span["text"] + " "
            metadata = {
                "source": pdf_path,
                "page": page_num + 1,
                "last_header": last_header
            }
            documents.append(Document(page_content=page_text, metadata=metadata))
    return documents

def extract_text_from_docx(docx_path: str) -> List[Document]:
    import docx
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    metadata = {
        "source": docx_path
    }
    return [Document(page_content=text, metadata=metadata)]

def extract_text_from_txt(txt_path: str) -> List[Document]:
    with open(txt_path, "r", encoding="utf-8") as file:
        text = file.read()
    metadata = {
        "source": txt_path
    }
    return [Document(page_content=text, metadata=metadata)]

@cl.on_chat_start
async def start():
    settings = await cl.ChatSettings(
        [
            Select(
                id="model_name",
                label="OpenAI - Model",
                values=["gpt-3.5-turbo", "gpt-3.5-turbo-16k", "gpt-4", "gpt-4-32k", "gpt-4o-mini"],
                initial_index=4,
            ),
            Slider(
                id="Chunk_Size",
                label="Chunk size",
                initial=600,
                min=100,
                max=3000,
                step=100,
            ),
            Slider(
                id="Max_Tokens",
                label="Max Tokens to use",
                initial=3000,
                min=100,
                max=20000,
                step=100,
            ),
            Slider(
                id="Font_Size_Threshold",
                label="Font Size Threshold",
                initial=11.9,
                min=5.0,
                max=30.0,
                step=0.1,
            ),
            Slider(
                id="Temperature",
                label="Temperature",
                initial=0.7,
                min=0,
                max=2,
                step=0.1,
            ),
            Slider(
                id="Chunk_Overlap",
                label="Chunk Overlap",
                initial=100,
                min=0,
                max=500,
                step=50,
            ),
            Slider(
                id="Number_Of_Chunks",
                label="Number of chunks to use",
                initial=3,
                min=1,
                max=20,
                step=1,
            ),
        ]
    ).send()
    await setup_agent(settings)

    await cl.Message(content="First send a message then upload files to begin.").send()

@cl.on_settings_update
async def setup_agent(settings):
    sentence_aware_splitter = SentenceAwareTextSplitter(
        chunk_size=settings["Chunk_Size"],
        chunk_overlap=settings["Chunk_Overlap"]
    )

    llm = ChatOpenAI(
        temperature=settings["Temperature"],
        model=settings["model_name"],
        max_tokens=settings["Max_Tokens"],
    )
    num_chunks = settings.get("Number_Of_Chunks", 5)

    cl.user_session.set("recursive_text_splitter", sentence_aware_splitter)
    cl.user_session.set("llm", llm)
    cl.user_session.set("num_chunks", num_chunks)
    font_size_threshold = settings.get("Font_Size_Threshold", 11.9)
    cl.user_session.set("Font_Size_Threshold", font_size_threshold)

    existing_chain = cl.user_session.get("chain")

    if existing_chain:
        chain = ConversationalRetrievalChain.from_llm(
            llm,
            chain_type="stuff",
            retriever=existing_chain.retriever,
            memory=existing_chain.memory,
            return_source_documents=True,
        )
        cl.user_session.set("chain", chain)
    else:
        print("Chain does not exist. Cannot update settings without an existing chain.")

@cl.on_message
async def main(message: cl.Message):
    if cl.user_session.get("chain") is None:
        await cl.Message(content="Please upload files to begin!").send()
        files = await cl.AskFileMessage(
            content="Upload files to proceed.",
            accept=["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain", "text/markdown"],
            max_size_mb=500,
            timeout=180,
            max_files=10
        ).send()

        all_texts = []

        for file in files:
            file_extension = os.path.splitext(file.name)[1].lower()
            msg = cl.Message(content=f"Processing {file.name}...")
            await msg.send()

            if file_extension == ".pdf":
                documents = extract_text_from_pdf(file.path)
            elif file_extension == ".docx":
                documents = extract_text_from_docx(file.path)
            elif file_extension == ".txt":
                documents = extract_text_from_txt(file.path)
            elif file_extension == ".md":
                with open(file.path, "r", encoding="utf-8") as file:
                    text = file.read()
                documents = [Document(page_content=text, metadata={"source": file.path})]
            else:
                msg.content = "Unsupported file type. Skipping this file."
                await msg.update()
                continue

            splitter = cl.user_session.get("markdown_text_splitter") if file_extension == ".md" else cl.user_session.get("recursive_text_splitter")

            for document in documents:
                chunks = splitter.split_text(document.page_content)
                for chunk in chunks:
                    chunk_metadata = document.metadata.copy()
                    chunk_metadata["chunk"] = chunk
                    all_texts.append(Document(page_content=chunk, metadata=chunk_metadata))

        filtered_documents = filter_complex_metadata(all_texts)

        embeddings = OpenAIEmbeddings()
        docsearch = await cl.make_async(Chroma.from_documents)(
            filtered_documents, embeddings
        )

        message_history = ChatMessageHistory()

        memory = ConversationBufferMemory(
            memory_key="chat_history",
            output_key="answer",
            chat_memory=message_history,
            return_messages=True,
        )

        chain = ConversationalRetrievalChain.from_llm(
            cl.user_session.get("llm"),
            chain_type="stuff",
            retriever=docsearch.as_retriever(),
            memory=memory,
            return_source_documents=True,
        )

        cl.user_session.set("chain", chain)

        await cl.Message(content="Processing done. You can now ask questions!").send()

    chain = cl.user_session.get("chain")
    cb = cl.AsyncLangchainCallbackHandler()

    res = await chain.acall(message.content, callbacks=[cb])
    answer = res["answer"]
    source_documents = res["source_documents"]

    num_chunks = cl.user_session.get("num_chunks", 3)
    
    if not isinstance(num_chunks, int):
        num_chunks = int(num_chunks)

    limited_documents = source_documents[:num_chunks]

    text_elements = []

    if limited_documents:
        for source_idx, source_doc in enumerate(limited_documents):
            source_name = f"source_{source_idx}"
            
            header = source_doc.metadata.get("last_header", "No header found")
            page = source_doc.metadata.get("page", "Unknown")

            content = f"### Page {page}\n\n{source_doc.page_content}\n\n---\n**Last Header:** {header}"
            
            text_elements.append(
                cl.Text(content=content, name=source_name, display="side")
            )
        
        prompt_length = len(message.content.split())
        response_length = len(answer.split())
        total_tokens = prompt_length + response_length
        token_usage_info = f"Estimated Token Usage: {total_tokens}"

        source_names = [text_el.name for text_el in text_elements]
        if source_names:
            answer += f"\n\n**Sources:** {', '.join(source_names)}"
        else:
            answer += "\n\n**No sources found**"
        
        answer += f"\n\n**{token_usage_info}**"
    else:
        answer += "\n\n**No documents found**"

    await cl.Message(content=answer, elements=text_elements).send()